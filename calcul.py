import streamlit as st
import pandas as pd
import numpy as np
import sqlite3
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import json

# Database setup
def init_db():
    conn = sqlite3.connect("tank.db")
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS calculations (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            timestamp TEXT,
            input_data TEXT,
            output_data TEXT
        )
    """)
    conn.commit()
    conn.close()

# Densitate function (ported from C++)
def densitate(T):
    d = 1 - (((T + 288.94) / (508929.2 * (T + 68.12963))) * (T - 3.9863)*(T - 3.9863))
    return d

# Processing logic

def process_data(df):
    w = {}     
    v = []     # Volumes
    m = []     # Masses
    i = 0;
    k = 0;
    M = 0;
    Trez1 = 0;
    v.append(0);
    hii = 0;
    for index, row in df.iterrows():
        Vi, Tvas, Trez, H = row['Vi'], row['Tvas'], row['Trez'], row['H']
        Vi = Vi * (1 + (0.00005) * (Tvas - 20))
        m.append(v[k] * densitate(Tvas))
        M = 0
        M = sum(m)
        k=k+1
        v.append(Vi)
        while(i <= H/10) :
            suma = 0;
            suma = M/densitate(Trez1);
            suma = suma * (1 + (2/3 * 0.000033 * (20 - Trez1)));
            VH = suma + Vi* ((i-hii)/(H/10 - hii))
            w[i] = int(VH)
            i = i + 1
        Trez1 = Trez
        hii = H/10
        T = Tvas
    
    # Format final results
    output_df = pd.DataFrame(sorted(w.items()), columns=['H (cm)', 'V (litri)'])
    return output_df

# Save to SQLite
def save_to_db(input_df, output_df):
    conn = sqlite3.connect("tank.db")
    cursor = conn.cursor()
    timestamp = datetime.now().isoformat()
    cursor.execute("INSERT INTO calculations (timestamp, input_data, output_data) VALUES (?, ?, ?)",
                   (timestamp, input_df.to_json(), output_df.to_json()))
    conn.commit()
    conn.close()

# Load past calculations from DB
def load_past_entries():
    conn = sqlite3.connect("tank.db")
    df = pd.read_sql_query("SELECT * FROM calculations ORDER BY timestamp DESC", conn)
    conn.close()
    return df

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for line in ["top", "bottom", "left", "right"]:
        element = OxmlElement(f"w:{line}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        tcPr.append(element)

def export_input(df):
    from math import ceil
    doc = Document()
    doc.add_heading('Rezultate Calcul Calibrare Rezervoare', level=1)
    num_cols = 5;
    max_rows_per_page = 25
    total_chunks = ceil(len(df) / (num_cols * max_rows_per_page))

    for chunk_index in range(total_chunks):
        chunk_start = chunk_index * (num_cols * max_rows_per_page)
        chunk_end = min(chunk_start + (num_cols * max_rows_per_page), len(df))
        chunk_df = df.iloc[chunk_start:chunk_end]

        total_rows = (len(chunk_df) + num_cols - 1) // num_cols
        table = doc.add_table(rows=total_rows+1, cols=num_cols*2)
        table.autofit = True;

        hdr_cells = table.rows[0].cells
        hdr_cells[1].text = 'Nr'
        hdr_cells[2].text = 'Vi'
        hdr_cells[2].text = 'Tvas'
        hdr_cells[2].text = 'Trez'
        hdr_cells[2].text = 'H'
        for i in range(num_cols):
            set_cell_border(hdr_cells[i])
        for r in range(total_rows):
            rows_cells = table.rows[r + 1].cells
            for c in range(num_cols):
                idx = (chunk_start) + (r + c * total_rows)
                if idx < len(df):
                    rows_cells[1].text = r + 1
                    rows_cells[2].text = str(df.iloc[idx, 1])
                    rows_cells[3].text = str(df.iloc[idx, 2])
                    rows_cells[4].text = str(df.iloc[idx, 3])
                    rows_cells[5].text = str(df.iloc[idx, 4])
                set_cell_border(rows_cells[c])
        doc.add_page_break()
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Export to Word document
def export_to_word(df):
    from math import ceil
    doc = Document()
    doc.add_heading('Rezultate Calcul Calibrare Rezervoare', level=1)
    num_cols = 6;
    max_rows_per_page = 25
    total_chunks = ceil(len(df) / (num_cols * max_rows_per_page))

    for chunk_index in range(total_chunks):
        chunk_start = chunk_index * (num_cols * max_rows_per_page)
        chunk_end = min(chunk_start + (num_cols * max_rows_per_page), len(df))
        chunk_df = df.iloc[chunk_start:chunk_end]

        total_rows = (len(chunk_df) + num_cols - 1) // num_cols
        table = doc.add_table(rows=total_rows+1, cols=num_cols*2)
        table.autofit = True;

        hdr_cells = table.rows[0].cells
        for i in range(num_cols):
            hdr_cells[i*2].text = 'H (cm)'
            hdr_cells[i*2 + 1].text = 'V (litri)'
            set_cell_border(hdr_cells[i*2])
            set_cell_border(hdr_cells[i*2 + 1])
        for r in range(total_rows):
            rows_cells = table.rows[r + 1].cells
            for c in range(num_cols):
                idx = (chunk_start) + (r + c * total_rows)
                if idx < len(df):
                    rows_cells[c*2].text = str(df.iloc[idx, 0])
                    rows_cells[c*2 + 1].text = str(df.iloc[idx, 1])
                else:
                    rows_cells[c*2].text = ''
                    rows_cells[c*2 + 1].text = ''
                set_cell_border(rows_cells[c*2])
                set_cell_border(rows_cells[c*2 + 1])
        doc.add_page_break()
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# Streamlit UI

init_db()
# Tabs for Calculator and Dashboard


st.title("Calcul Calibrare Rezervoare")
st.write("Introdu datele mai jos:")

sample_data = pd.DataFrame({
    'Vi': [0],
    'Tvas': [20],
    'Trez': [20],
    'H': [0]
})

input_df = st.data_editor(sample_data, num_rows="dynamic", use_container_width=True)

if 'output_df' not in st.session_state:
    st.session_state.output_df = None
    st.session_state.doc_file = None
    st.session_state.filename = None

if st.button("Calculate"):
    try:
        output_df = process_data(input_df)
        doc_file = export_to_word(output_df)
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        filename = f"tank_volume_{timestamp}.docx"

        # Save to session state
        st.session_state.output_df = output_df
        st.session_state.doc_file = doc_file
        st.session_state.filename = filename

        # Save to DB only once
        save_to_db(input_df, output_df)

        st.success("Calculation successful!")
    except Exception as e:
        st.error(f"Error during calculation: {e}")

# # Show output if available
# if st.session_state.output_df is not None:
#     st.dataframe(st.session_state.output_df, use_container_width=True)

#     st.download_button(
#         label="Download as Word Document",
#         data=st.session_state.doc_file,
#         file_name=st.session_state.filename,
#         mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#     )

st.header("Istoric")
past_df = load_past_entries()
if past_df.empty:
    st.info("Nu avem nimic in istoric.")
else:
    date_list = [datetime.fromisoformat(ts).strftime('%d/%m/%Y, %H:%M:%S') for ts in past_df['timestamp']]
    selected_timestamp = st.selectbox("Selecteaza o data:", date_list)
    selected_row = past_df[past_df['timestamp'].apply(lambda ts: datetime.fromisoformat(ts).strftime('%d/%m/%Y, %H:%M:%S')) == selected_timestamp].iloc[0]

    st.subheader("Input Data")
    st.dataframe(pd.read_json(selected_row['input_data']), use_container_width=True)

    st.subheader("Output Data")
    st.dataframe(pd.read_json(selected_row['output_data']), use_container_width=True)
    
    selected_output_df = pd.read_json(selected_row['output_data'])
    selected_doc_file = export_to_word(selected_output_df)

    download_filename = f"tank_volume_{datetime.fromisoformat(selected_row['timestamp']).strftime('%Y-%m-%d_%H-%M-%S')}.docx"

    st.download_button(
        label="Download Selected Output as Word Document",
        data=selected_doc_file,
        file_name=download_filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

